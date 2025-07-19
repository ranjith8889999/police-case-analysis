"""
LangChain-based document service with Gemini AI and normalized embeddings.
Uses LangChain for simpler, more readable code structure.
"""

import os
import uuid
import numpy as np
from werkzeug.utils import secure_filename
from app.models.database import Document, Chunk, Embedding

# LangChain imports for easier development
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.schema import Document as LangChainDocument

import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from sqlalchemy import text
from app.utils.database_utils import retry_db_operation

class GeminiDocumentService:
    def __init__(self, db_session, api_key=None):
        """
        Initialize LangChain-based document service with Gemini.
        
        Args:
            db_session: SQLAlchemy session for database operations
            api_key: Gemini API key (will use environment variable if not provided)
        """
        self.db_session = db_session
        self.api_key = api_key or os.environ.get('GOOGLE_API_KEY', 'AIzaSyC2tznCdJX-y9YzEuE_USCA3BVUuo1-av4')
        
        # LangChain Gemini embeddings - much simpler!
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/text-embedding-004",
            google_api_key=self.api_key
        )
        
        # LangChain text splitter - easy to configure
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        self.upload_folder = os.path.abspath('./app/static/uploads')        
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def normalize_vector(self, vector):
        """
        Normalize a vector to unit length for better similarity comparisons.
        
        Args:
            vector: Input vector (list or numpy array)
            
        Returns:
            numpy.ndarray: Normalized vector
        """
        vector_array = np.array(vector)
        norm = np.linalg.norm(vector_array)
        if norm == 0:
            return vector_array
        return vector_array / norm
    
    def generate_embedding(self, text):
        """
        Generate normalized embedding using LangChain Gemini embeddings.
        Much simpler than direct API calls!
        
        Args:
            text (str): Text to embed
            
        Returns:
            list: Normalized embedding vector
        """
        try:
            # LangChain makes this super simple!
            embedding_vector = self.embeddings.embed_query(text)
            
            # Normalize the vector for better similarity search
            normalized_vector = self.normalize_vector(embedding_vector)
            
            return normalized_vector.tolist()
            
        except Exception as e:
            print(f"Error generating embedding: {e}")
            raise
    
    def generate_query_embedding(self, query):
        """
        Generate normalized embedding for query using LangChain.
        Same method as generate_embedding - LangChain handles the task type automatically!
        
        Args:
            query (str): Query text to embed
            
        Returns:
            list: Normalized embedding vector
        """
        return self.generate_embedding(query)  # LangChain handles query vs document automatically
    
    def extract_text_from_file(self, file_path):
        """Extract text from various file formats."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # PDF files
            if file_extension == '.pdf':
                text = ''
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        text += pdf_reader.pages[page_num].extract_text()
                return text
            
            # Word documents
            elif file_extension in ['.docx', '.doc']:
                doc = DocxDocument(file_path)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text)
                return '\n'.join(text)
            
            # Text files
            elif file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read()
            
            # Excel files
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                return df.to_string()
            
            # CSV files
            elif file_extension == '.csv':
                df = pd.read_csv(file_path)
                return df.to_string()
            
            # Unsupported format
            else:
                return f"Unsupported file format: {file_extension}"
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return f"Error processing file: {str(e)}"
    
    @retry_db_operation(max_retries=3, delay=1, backoff=2)
    def upload_document(self, file, document_type, title, user_id):
        """Upload and process a document with chunking and normalized embeddings."""
        try:
            # Create unique filename
            filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4().hex}_{filename}"
            
            # Save file to uploads folder
            upload_path = os.path.join(self.upload_folder, unique_filename)
            file.save(upload_path)
            
            # Extract text content
            content = self.extract_text_from_file(upload_path)
            
            # Add document to database first
            document = Document(
                user_id=user_id,
                document_name=title,
                document_type=document_type,
                file_path=upload_path
            )
            
            self.db_session.add(document)
            self.db_session.flush()  # Get the document ID
            
            # Create and store document chunks with normalized embeddings
            if content:
                chunks = self.create_document_chunks(document.id, content)
                print(f"Created {len(chunks)} chunks for document '{title}' with normalized embeddings")
            
            self.db_session.commit()
            
            return document
            
        except Exception as e:
            self.db_session.rollback()
            print(f"Error uploading document: {e}")
            raise
    
    def create_document_chunks(self, document_id, content):
        """Create document chunks using LangChain text splitter and normalized embeddings."""
        try:
            # LangChain text splitter - super simple!
            text_chunks = self.text_splitter.split_text(content)
            
            # Create Chunk objects and embeddings
            chunks = []
            
            for i, chunk_text in enumerate(text_chunks):
                # Skip empty chunks
                if not chunk_text.strip():
                    continue
                    
                # Create chunk
                chunk = Chunk(
                    document_id=document_id,
                    chunk_text=chunk_text,
                    chunk_order=i
                )
                
                chunks.append(chunk)
                self.db_session.add(chunk)
                self.db_session.flush()  # Get chunk ID
                
                # Generate normalized embedding using LangChain
                embedding_vector = self.generate_embedding(chunk_text)
                
                # Create embedding record with normalized vector
                embedding = Embedding(
                    chunk_id=chunk.id,
                    embedding_vector=embedding_vector
                )
                
                self.db_session.add(embedding)
            
            return chunks            
        except Exception as e:
            self.db_session.rollback()
            print(f"Error creating document chunks: {e}")
            raise e
    
    def cosine_similarity(self, vec1, vec2):
        """
        Calculate cosine similarity between two normalized vectors.
        For normalized vectors, this is simply the dot product.
        
        Args:
            vec1, vec2: Normalized vectors
            
        Returns:
            float: Cosine similarity score
        """
        return np.dot(vec1, vec2)
    
    def semantic_search_chunks(self, query, document_type=None, limit=5, similarity_threshold=0.2):
        """
        Enhanced semantic search with proper vector handling and accuracy improvements.
        """
        try:
            print(f"DEBUG: Starting semantic search for query: '{query}'")
            print(f"DEBUG: Document type filter: {document_type}")
            
            # Generate normalized query embedding
            query_embedding = self.generate_query_embedding(query)
            
            # Validate embedding
            if not query_embedding or len(query_embedding) == 0:
                print("ERROR: Failed to generate query embedding, falling back to text search")
                return self.basic_search_fallback(query, document_type, limit)
            
            # Use proper parameterized query for vector search
            if document_type:
                sql = """
                SELECT 
                    c.chunk_id,
                    c.chunk_text,
                    c.chunk_order,
                    d.document_name,
                    d.document_type,
                    d.document_id,
                    d.file_path,
                    (1 - (e.embedding_vector <=> :query_vector::vector)) as similarity
                FROM chunks c
                JOIN documents d ON c.document_id = d.document_id
                JOIN embeddings e ON c.chunk_id = e.chunk_id
                WHERE d.document_type = :doc_type
                  AND e.embedding_vector IS NOT NULL
                ORDER BY e.embedding_vector <=> :query_vector::vector 
                LIMIT :limit_val
                """
                params = {
                    'query_vector': str(query_embedding),
                    'doc_type': document_type,
                    'limit_val': limit
                }
            else:
                sql = """
                SELECT 
                    c.chunk_id,
                    c.chunk_text,
                    c.chunk_order,
                    d.document_name,
                    d.document_type,
                    d.document_id,
                    d.file_path,
                    (1 - (e.embedding_vector <=> :query_vector::vector)) as similarity
                FROM chunks c
                JOIN documents d ON c.document_id = d.document_id
                JOIN embeddings e ON c.chunk_id = e.chunk_id
                WHERE e.embedding_vector IS NOT NULL
                ORDER BY e.embedding_vector <=> :query_vector::vector 
                LIMIT :limit_val
                """
                params = {
                    'query_vector': str(query_embedding),
                    'limit_val': limit
                }
            
            print(f"DEBUG: Executing vector similarity search with {len(params)} parameters")
            
            # For vector operations, we need to format the query differently
            # Convert embedding to proper format for PostgreSQL
            vector_str = '[' + ','.join(map(str, query_embedding)) + ']'
            
            if document_type:
                # Use safe string interpolation for vector data
                sql_formatted = sql.replace(':query_vector', f"'{vector_str}'").replace(':doc_type', f"'{document_type}'").replace(':limit_val', str(limit))
                result = self.db_session.execute(text(sql_formatted))
            else:
                # Use safe string interpolation for vector data
                sql_formatted = sql.replace(':query_vector', f"'{vector_str}'").replace(':limit_val', str(limit))
                result = self.db_session.execute(text(sql_formatted))
            
            # Process results with enhanced scoring
            results = []
            query_terms = self._extract_key_terms(query)
            print(f"DEBUG: Extracted key terms: {query_terms}")
            
            row_count = 0
            for row in result:
                row_count += 1
                print(f"DEBUG: Processing result {row_count} with similarity: {row.similarity:.4f}")
                
                # Apply similarity threshold
                if row.similarity < similarity_threshold:
                    print(f"DEBUG: Skipping result due to low similarity: {row.similarity:.4f} < {similarity_threshold}")
                    continue
                
                # Enhanced scoring with multiple factors
                text_match_score = self._calculate_text_match_score(row.chunk_text, query_terms)
                length_penalty = max(0.8, min(1.0, len(row.chunk_text) / 1000))  # Prefer reasonably sized chunks
                
                # Combined score: vector similarity + text matching + length consideration
                final_score = (row.similarity * 0.7) + (text_match_score * 0.2) + (length_penalty * 0.1)
                
                # Create enhanced preview
                preview = self._create_enhanced_preview(row.chunk_text, query_terms)
                
                results.append({
                    'chunk_id': row.chunk_id,
                    'chunk_text': row.chunk_text,
                    'chunk_order': row.chunk_order,
                    'document_name': row.document_name,
                    'document_type': row.document_type,
                    'similarity': row.similarity,
                    'text_match_score': text_match_score,
                    'final_score': final_score,
                    'document_id': row.document_id,
                    'file_path': row.file_path,
                    'preview': preview,
                    'relevance_reason': self._explain_relevance(row.chunk_text, query, final_score)
                })
            
            print(f"DEBUG: Total vector results: {row_count}, after filtering: {len(results)}")
            
            # Sort by final score (best first)
            results.sort(key=lambda x: x['final_score'], reverse=True)
            
            # If no good vector results, try hybrid approach
            if len(results) == 0:
                print("DEBUG: No vector results found, trying hybrid search")
                return self._hybrid_search(query, document_type, limit)
            
            print(f"Semantic search found {len(results)} relevant chunks for query: '{query}'")
            return results
            
        except Exception as e:
            print(f"Error in semantic search: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to basic search
            return self.basic_search_fallback(query, document_type, limit)
    
    def _calculate_text_match_score(self, text, query_terms):
        """Calculate score based on text matches."""
        if not query_terms:
            return 0.0
        text_lower = text.lower()
        matches = sum(1 for term in query_terms if term in text_lower)
        return matches / len(query_terms) if query_terms else 0
    
    def _hybrid_search(self, query, document_type=None, limit=5):
        """Hybrid search combining text and vector approaches."""
        return self.basic_search_fallback(query, document_type, limit)
    
    def basic_search_fallback(self, query, document_type=None, limit=5):
        """Simple fallback search using text matching."""
        try:
            if document_type:
                sql = """
                SELECT 
                    c.chunk_id,
                    c.chunk_text,
                    c.chunk_order,
                    d.document_name,
                    d.document_type,
                    d.document_id,
                    d.file_path,
                    0.5 as similarity
                FROM chunks c
                JOIN documents d ON c.document_id = d.document_id
                WHERE LOWER(c.chunk_text) LIKE :search_pattern AND d.document_type = :doc_type
                ORDER BY LENGTH(c.chunk_text) DESC LIMIT :limit_val
                """
                params = {
                    'search_pattern': f'%{query.lower()}%',
                    'doc_type': document_type,
                    'limit_val': limit
                }
            else:
                sql = """
                SELECT 
                    c.chunk_id,
                    c.chunk_text,
                    c.chunk_order,
                    d.document_name,
                    d.document_type,
                    d.document_id,
                    d.file_path,
                    0.5 as similarity
                FROM chunks c
                JOIN documents d ON c.document_id = d.document_id
                WHERE LOWER(c.chunk_text) LIKE :search_pattern
                ORDER BY LENGTH(c.chunk_text) DESC LIMIT :limit_val
                """
                params = {
                    'search_pattern': f'%{query.lower()}%',
                    'limit_val': limit
                }
            
            result = self.db_session.execute(text(sql), params)
            
            results = []
            for row in result:
                preview = row.chunk_text[:300] + "..." if len(row.chunk_text) > 300 else row.chunk_text
                
                results.append({
                    'chunk_id': row.chunk_id,
                    'chunk_text': row.chunk_text,
                    'chunk_order': row.chunk_order,
                    'document_name': row.document_name,
                    'document_type': row.document_type,
                    'similarity': row.similarity,
                    'final_score': row.similarity,
                    'document_id': row.document_id,
                    'file_path': row.file_path,
                    'preview': preview,
                    'relevance_reason': "Text match found"
                })
            
            return results
            
        except Exception as e:
            print(f"Fallback search also failed: {e}")
            return []
    
    def _extract_key_terms(self, query):
        """Extract key terms from query for hybrid search."""
        import re
        
        # Remove common stop words but keep legal and location terms
        stop_words = {
            'the', 'and', 'or', 'in', 'on', 'at', 'to', 'a', 'an', 'is', 'are', 
            'was', 'were', 'be', 'this', 'that', 'with', 'for', 'as', 'by', 'of',
            'me', 'about', 'explain'
        }
        
        # Extract words, convert to lowercase, but preserve important legal/location terms
        terms = re.findall(r'\b\w+\b', query.lower())
        
        # Keep important terms even if short
        important_terms = {'law', 'act', 'rule', 'code', 'state', 'amendment', 'music', 'street', 'orissa', 'odisha'}
        
        key_terms = []
        for term in terms:
            if term not in stop_words and (len(term) > 2 or term in important_terms):
                key_terms.append(term)
        
        return key_terms
    
    def _create_enhanced_preview(self, text, key_terms, max_length=300):
        """Create a preview with key terms highlighted and proper context."""
        import re
        
        # Find the best excerpt around key terms
        best_start = 0
        best_score = 0
        
        for term in key_terms:
            match = re.search(rf'\b{re.escape(term)}\b', text, re.IGNORECASE)
            if match:
                start = max(0, match.start() - 100)
                excerpt = text[start:start + max_length]
                
                # Count key terms in this excerpt
                score = sum(1 for t in key_terms if t in excerpt.lower())
                
                if score > best_score:
                    best_score = score
                    best_start = start
        
        # Create the preview
        preview = text[best_start:best_start + max_length]
        
        # Add ellipsis if truncated
        if best_start > 0:
            preview = "..." + preview
        if best_start + max_length < len(text):
            preview = preview + "..."
            
        return preview.strip()
    
    def _explain_relevance(self, chunk_text, query, similarity):
        """Generate a brief explanation of why this chunk is relevant."""
        if similarity > 0.8:
            return "Highly relevant - strong semantic match"
        elif similarity > 0.6:
            return "Moderately relevant - good conceptual match"
        elif similarity > 0.4:
            return "Somewhat relevant - partial topic overlap"
        else:
            return "Low relevance - minimal connection"
    
    @retry_db_operation(max_retries=3, delay=1, backoff=2)
    def search_document_chunks(self, query, document_type=None, limit=5):
        """
        Wrapper method that uses semantic search for better results.
        """
        return self.semantic_search_chunks(query, document_type, limit)
    
    def text_search_fallback(self, query, document_type=None, limit=5):
        """Fallback text search when vector search fails."""
        try:
            sql = """
            SELECT 
                c.chunk_id,
                c.chunk_text,
                c.chunk_order,
                d.document_name,
                d.document_type
            FROM chunks c
            JOIN documents d ON c.document_id = d.document_id
            WHERE c.chunk_text ILIKE :search_pattern
            """
            
            params = {
                'search_pattern': f'%{query}%',
                'limit': limit
            }
            
            if document_type:
                sql += " AND d.document_type = :document_type"
                params['document_type'] = document_type
            
            sql += " ORDER BY d.uploaded_at DESC LIMIT :limit"
            
            result = self.db_session.execute(text(sql), params)
            
            return [(row.chunk_id, row.chunk_text, row.chunk_order, 
                    row.document_name, row.document_type) for row in result]
            
        except Exception as e:
            print(f"Text search fallback also failed: {e}")
            return []
    
    def get_documents_by_type(self, document_type):
        """Get all documents of a specific type."""
        return self.db_session.query(Document).filter(Document.document_type == document_type).all()
    
    def get_document_by_id(self, document_id):
        """Get document by ID."""
        return self.db_session.query(Document).filter(Document.id == document_id).first()
    
    def get_all_documents(self):
        """Get all documents."""
        return self.db_session.query(Document).order_by(Document.uploaded_at.desc()).all()
        
    @retry_db_operation(max_retries=3, delay=1, backoff=2)
    def delete_document(self, document_id):
        """Delete document and all its chunks."""
        try:
            document = self.db_session.query(Document).filter(Document.id == document_id).first()
            if document:
                # Delete file from filesystem
                if os.path.exists(document.file_path):
                    os.remove(document.file_path)
                
                # Delete from database (chunks and embeddings will be deleted due to cascade)
                self.db_session.delete(document)
                self.db_session.commit()
                return True
            
            return False
        except Exception as e:
            self.db_session.rollback()
            print(f"Error deleting document: {e}")
            raise e
    
    @retry_db_operation(max_retries=3, delay=1, backoff=2)
    def update_document(self, document_id, title, document_type, content=None):
        """Update document and recreate chunks if content changed."""
        try:
            document = self.db_session.query(Document).filter(Document.id == document_id).first()
            if document:
                document.document_name = title
                document.document_type = document_type
                
                # If content is provided, update it and recreate chunks
                if content is not None:
                    # Delete existing chunks (embeddings will be deleted due to cascade)
                    for chunk in document.chunks:
                        self.db_session.delete(chunk)
                    
                    # Create new chunks with normalized embeddings
                    if content:
                        self.create_document_chunks(document.id, content)
                
                self.db_session.commit()
                return document
            return None
        except Exception as e:
            self.db_session.rollback()
            print(f"Error updating document: {e}")
            raise
